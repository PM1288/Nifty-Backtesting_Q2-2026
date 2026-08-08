#!/usr/bin/env python3
"""Universal V2 fail-closed evaluator for a completed strategy run folder."""
from __future__ import annotations
import argparse, hashlib, json, math, os, shutil, uuid, zipfile
from datetime import datetime, timezone
from pathlib import Path
import pandas as pd

ROOT=Path(__file__).resolve().parents[3]
POLICY=ROOT/'platform/nifty_stratlab/config/evaluation/universal_strategy_evaluation_v2.json'
DDL=ROOT/'db/sql/028_universal_strategy_evaluation.sql'
SHEETS=['00 Executive Dashboard','01 Strategy Map','02 Validation Gates','03 Signal Funnel','04 Reward Ladder','05 Adverse Ladder','06 Target vs Adverse Sequence','07 Trade Quality','08 D+5 Outcome','09 30-Session Diagnostic','10 Finite-Capital Portfolio','11 Daily Equity and Drawdown','12 Annual Performance','13 Monthly Performance','14 Stock Performance','15 Sector Performance','16 Regime and P-Diagram','17 Factor-Rule Importance','18 Open-Trapped Positions','19 Skipped Signals','20 Data Coverage-Quality','21 Risk Register','22 Trade Detail','23 Assumptions-Versions','24 Evidence Index']
def sha(p):
 h=hashlib.sha256()
 with p.open('rb') as f:
  for b in iter(lambda:f.read(1024*1024),b''): h.update(b)
 return h.hexdigest()
def load(folder,name):
 p=folder/name
 return pd.read_csv(p,low_memory=False) if p.exists() else pd.DataFrame()
def number(s): return pd.to_numeric(s,errors='coerce')
def pct(flag):
 s=pd.Series(flag).dropna(); return round(float(s.astype(bool).mean()*100),4) if len(s) else None
def metric(k,v,unit='',status='ESTIMATED',note=''): return {'metric':k,'value':v,'unit':unit,'status':status,'note':note}
def risk_rows():
 vals=[('R01','Look-ahead or same-candle sequencing','HIGH','MEDIUM','target timing','Enforce post-entry timestamps and flag ambiguity','MEDIUM','CONDITIONAL'),('R02','Survivorship/current-universe bias','HIGH','HIGH','generalisability','Use point-in-time constituent snapshots','HIGH','CONDITIONAL'),('R03','Open losses excluded from success metrics','CRITICAL','MEDIUM','P&L and win rate','Include open MTM liability and separate closed/open','MEDIUM','YES'),('R04','MFE confused with realised return','CRITICAL','MEDIUM','return','Label opportunity separately from replay','LOW','YES'),('R05','Missing authoritative exit','CRITICAL','MEDIUM','realised P&L','Fail closed unless exit source is supplied','HIGH','YES'),('R06','No finite-capital chronology','CRITICAL','MEDIUM','portfolio return','Require allocator/equity evidence','HIGH','YES'),('R07','Overlapping labels/trades','HIGH','HIGH','confidence','Use purged chronological folds and effective sample size','MEDIUM','CONDITIONAL'),('R08','Costs, spread and slippage uncertainty','HIGH','MEDIUM','net economics','Version cost model and run sensitivities','MEDIUM','CONDITIONAL'),('R09','End-of-data right censoring','HIGH','MEDIUM','D30 and recovery','Retain and exclude censored outcomes from mature denominators','LOW','CONDITIONAL'),('R10','Multiple testing/parameter selection','HIGH','HIGH','robustness','Immutable trial ledger and adjusted evidence','HIGH','CONDITIONAL')]
 return pd.DataFrame(vals,columns=['risk_id','description','severity','probability','affected_metric','mitigation','residual_risk','invalidates_conclusion'])
def gate_rows(archetype, trades, h30, targets, config, authoritative_exit=False, equity=None, trials=None):
 exit_known=authoritative_exit and ('exit_reason' in trades and trades.exit_reason.notna().any())
 finite=equity is not None and len(equity)>1
 oos=trials is not None and len(trials)>0 and trials.astype(str).apply(lambda c:c.str.contains('OUT_OF_SAMPLE|OOS',case=False,regex=True)).any().any()
 return pd.DataFrame([
  ('data_quality','PASS' if len(trades) else 'FAIL',f'{len(trades)} trade rows; coverage flags retained'),
  ('temporal_integrity','WARN' if not len(targets) else 'PASS','Event timestamps available' if len(targets) else 'No target-event sequencing file'),
  ('authoritative_exit_for_realised_pnl','PASS' if exit_known else 'FAIL','Exit evidence present' if exit_known else 'Entry-only strategy: realised strategy P&L not estimable'),
  ('finite_capital_for_portfolio_return','PASS' if finite else 'FAIL','Capital release fields present' if finite else 'No complete portfolio chronology'),
  ('out_of_sample_evidence','PASS' if oos else 'FAIL','Chronological OOS trial evidence supplied' if oos else 'No immutable chronological OOS fold/trial ledger supplied to this evaluator'),
  ('reproducibility','PASS','Input hashes and artifact manifest generated'),
  ('h30_maturity','PASS' if len(h30) else 'WARN',f'{len(h30)} H30 observations'),
 ],columns=['gate_name','status','evidence'])
def main():
 p=argparse.ArgumentParser(); p.add_argument('--input-dir',type=Path,required=True); p.add_argument('--strategy-name',required=True); p.add_argument('--strategy-version',default='UNKNOWN'); p.add_argument('--archetype',choices=['ENTRY_ONLY','COMPLETE_RULE_BASED','SCORE_BASED','RANKING_MODEL','MULTI_STAGE','HYBRID','OPTIONS','OTHER'],required=True); p.add_argument('--evaluation-mode',default='PATH_AND_SHARED_ROE'); p.add_argument('--authoritative-exit',action='store_true',help='Assert that exit_reason comes from the strategy contract, not shared RoE'); p.add_argument('--output-dir',type=Path,required=True); p.add_argument('--database-url',default=os.getenv('DATABASE_URL') or os.getenv('TRADING_DATABASE_URL')); a=p.parse_args(); cfg=json.loads(POLICY.read_text()); run_id=str(uuid.uuid4()); out=a.output_dir; out.mkdir(parents=True,exist_ok=True)
 trades=load(a.input_dir,'trades.csv'); targets=load(a.input_dir,'target_events.csv'); adverse=load(a.input_dir,'adverse_events.csv'); h30=load(a.input_dir,'h30_observations.csv'); decisions=load(a.input_dir,'decisions.csv'); regimes=load(a.input_dir,'regime_performance.csv'); equity=load(a.input_dir,'portfolio_equity.csv'); portfolio=load(a.input_dir,'portfolio_decisions.csv'); trials=load(a.input_dir,'trial_ledger.csv'); skipped=load(a.input_dir,'skipped_signals.csv')
 if trades.empty: raise SystemExit(f'Missing or empty {a.input_dir}/trades.csv')
 trades.insert(0,'evaluation_run_id',run_id); trades.insert(1,'strategy_name',a.strategy_name); trades.insert(2,'strategy_version',a.strategy_version); trades.insert(3,'evaluation_mode',a.evaluation_mode)
 key='entry_path_id' if 'entry_path_id' in trades else None
 if key and len(h30):
  keep=[c for c in [key,'maturity_status','rankable_flag','coverage_status','return_d29_pct','mae_30t_pct','underwater_session_count','longest_underwater_streak','max_close_return_so_far_pct','maximum_high_30t','minimum_low_30t'] if c in h30]
  trades=trades.merge(h30[keep].drop_duplicates(key),on=key,how='left',suffixes=('','_h30'))
 for c in ['mfe_pct','mae_pct','after_tax_net_pnl','unrealized_net_liquidation_pnl','gross_pnl','costs','holding_sessions','return_d29_pct','mae_30t_pct']: 
  if c in trades: trades[c]=number(trades[c])
 if 'entry_date' in trades: trades['entry_date']=pd.to_datetime(trades.entry_date,errors='coerce'); trades['entry_year']=trades.entry_date.dt.year; trades['entry_month']=trades.entry_date.dt.to_period('M').astype(str)
 status=trades.get('status',pd.Series('',index=trades.index)).astype(str).str.upper(); trades['is_open']=status.isin(['OPEN','UNRESOLVED_OPEN'])
 trades['net_result']=number(trades.get('after_tax_net_pnl',pd.Series(index=trades.index,dtype=float))).fillna(0)+number(trades.get('unrealized_net_liquidation_pnl',pd.Series(index=trades.index,dtype=float))).fillna(0)
 trades['net_profitable']=trades.net_result>0; trades['persistent_loser']=(number(trades.get('mae_pct',pd.Series(index=trades.index,dtype=float)))<=-5)&(~trades.net_profitable); trades['high_risk_eventual_winner']=(number(trades.get('mae_pct',pd.Series(index=trades.index,dtype=float)))<=-5)&trades.net_profitable
 trade_path=out/f'{a.strategy_name}_Trades.csv'; trades.to_csv(trade_path,index=False)
 gates=gate_rows(a.archetype,trades,h30,targets,cfg,a.authoritative_exit,equity,trials); risks=risk_rows(); scorable=not (gates.status=='FAIL').any(); state='EXPERIMENTAL' if scorable else ('NOT_SCORABLE_DATA_FAILURE' if gates.loc[gates.gate_name=='data_quality','status'].iloc[0]=='FAIL' else 'NOT_SCORABLE_METHOD_FAILURE')
 reward=[]
 if len(targets):
  hit=targets.get('hit_flag',targets.get('touched',False)); levels=targets.assign(_hit=pd.Series(hit).astype(str).str.lower().isin(['true','1','t','yes'])).groupby(['level_id','level_pct'],dropna=False)._hit.agg(['count','sum']).reset_index(); levels['hit_rate_pct']=levels['sum']/levels['count']*100; reward=levels
 adverse_summary=pd.DataFrame()
 if len(adverse):
  hit=adverse.get('hit_flag',adverse.get('touched',False)); adverse_summary=adverse.assign(_hit=pd.Series(hit).astype(str).str.lower().isin(['true','1','t','yes'])).groupby(['level_id','level_pct'],dropna=False)._hit.agg(['count','sum']).reset_index(); adverse_summary['breach_rate_pct']=adverse_summary['sum']/adverse_summary['count']*100
 summary=pd.DataFrame([metric('Validation state',state),metric('Overall score','NOT SCORABLE' if not scorable else None,status='BLOCKED' if not scorable else 'PENDING'),metric('Trades',len(trades),'count'),metric('Symbols',trades.symbol.nunique() if 'symbol' in trades else None,'count'),metric('Closed trades',int((~trades.is_open).sum()),'count'),metric('Open positions',int(trades.is_open.sum()),'count'),metric('Net profitable rate',pct(trades.net_profitable),'%'),metric('Median MFE',trades.mfe_pct.median() if 'mfe_pct' in trades else None,'%','DIAGNOSTIC','Not realised P&L'),metric('Median MAE',trades.mae_pct.median() if 'mae_pct' in trades else None,'%'),metric('Net P&L including open MTM',trades.net_result.sum(),'INR','ESTIMATED'),metric('H30 mature rows',int((h30.get('maturity_status',pd.Series(dtype=str)).astype(str)=='MATURE').sum()),'count')])
 funnel=pd.DataFrame([('Raw decision rows',len(decisions)),('Trade rows',len(trades)),('Closed trades',int((~trades.is_open).sum())),('Open trades',int(trades.is_open.sum()))],columns=['stage','count'])
 def grouped(keys):
  if not all(k in trades for k in keys): return pd.DataFrame({'status':['NOT ESTIMABLE: missing '+','.join(keys)]})
  return trades.groupby(keys,dropna=False).agg(trades=('net_result','size'),net_pnl=('net_result','sum'),median_net=('net_result','median'),profitable_rate_pct=('net_profitable',lambda s:s.mean()*100),open_positions=('is_open','sum'),median_mfe_pct=('mfe_pct','median') if 'mfe_pct' in trades else ('net_result','size'),median_mae_pct=('mae_pct','median') if 'mae_pct' in trades else ('net_result','size')).reset_index()
 annual=grouped(['entry_year']); monthly=grouped(['entry_month']); stocks=grouped(['symbol']); sectors=grouped(['sector'])
 open_trapped=trades[trades.is_open | trades.persistent_loser].copy(); coverage=pd.DataFrame([{'source_file':p.name,'rows':len(load(a.input_dir,p.name)),'sha256':sha(p)} for p in a.input_dir.glob('*.csv') if p.name in ['trades.csv','target_events.csv','adverse_events.csv','h30_observations.csv','decisions.csv','regime_performance.csv']])
 strategy_map=pd.DataFrame([('strategy_name',a.strategy_name),('strategy_version',a.strategy_version),('archetype',a.archetype),('evaluation_mode',a.evaluation_mode),('entry_authority','Preserved from source run'),('exit_authority','AUTHORITATIVE' if a.authoritative_exit else 'NOT ASSERTED: shared RoE is diagnostic/scenario only'),('policy_id',cfg['policy_id'])],columns=['field','value'])
 pdiag=pd.DataFrame([('INPUT','Signals, OHLCV, instrument, direction, entry anchor'),('SYSTEM','Strategy version, entry engine, exit engine, evaluator, portfolio allocator'),('NOISE','Nifty/stock regime, VIX, gaps, liquidity, events, corporate actions, stale data'),('CONTROL','Thresholds, timing, sizing, targets, carry, data gates, costs'),('ERROR','Wrong direction; target/time; economics/execution; data/risk/portfolio')],columns=['type','detail'])
 assumptions=pd.DataFrame([('Generated UTC',datetime.now(timezone.utc).isoformat()),('Input directory',str(a.input_dir.resolve())),('Policy file',str(POLICY)),('Authoritative exit asserted',a.authoritative_exit),('Realised P&L rule','Only authoritative replay; entry-only shared-RoE results remain scenario estimates'),('Portfolio return rule','Only valid with complete finite-capital chronology'),('Tax','Not treated as universal transaction cost; retain source tax reserve separately')],columns=['item','value'])
 empty=lambda reason:pd.DataFrame({'status':[reason]})
 sheets={SHEETS[0]:summary,SHEETS[1]:strategy_map,SHEETS[2]:gates,SHEETS[3]:funnel,SHEETS[4]:reward if len(reward) else empty('No target_events.csv'),SHEETS[5]:adverse_summary if len(adverse_summary) else empty('No adverse_events.csv'),SHEETS[6]:targets.groupby('sequence',dropna=False).size().reset_index(name='events') if 'sequence' in targets else empty('Sequence not available'),SHEETS[7]:grouped(['net_profitable','persistent_loser','high_risk_eventual_winner']),SHEETS[8]:reward[reward.level_pct.isin([1,2,5])] if len(reward) else empty('D+5 targets unavailable'),SHEETS[9]:h30 if len(h30) else empty('H30 unavailable'),SHEETS[10]:portfolio if len(portfolio) else empty('Use authoritative finite-capital portfolio ledger; not inferred from independent opportunities'),SHEETS[11]:equity if len(equity) else empty('Daily equity not supplied; portfolio drawdown not estimable'),SHEETS[12]:annual,SHEETS[13]:monthly,SHEETS[14]:stocks,SHEETS[15]:sectors,SHEETS[16]:pd.concat([pdiag,regimes.astype(str)],ignore_index=True) if len(regimes) else pdiag,SHEETS[17]:trials if len(trials) else empty('Component importance requires factor values and chronological OOS ablation; parameter sensitivity is separate'),SHEETS[18]:open_trapped,SHEETS[19]:skipped if len(skipped) else empty('Skipped-signal ledger not supplied'),SHEETS[20]:coverage,SHEETS[21]:risks,SHEETS[22]:trades,SHEETS[23]:assumptions,SHEETS[24]:empty('Populated after artifact generation')}
 xlsx=out/f'{a.strategy_name}_Evaluation_Results.xlsx'
 with pd.ExcelWriter(xlsx,engine='xlsxwriter') as w:
  for name,df in sheets.items():
   df.head(1048500).to_excel(w,sheet_name=name,index=False); ws=w.sheets[name]; ws.freeze_panes(1,0); ws.autofilter(0,0,max(len(df),1),max(len(df.columns)-1,0)); ws.set_column(0,max(len(df.columns)-1,0),18)
 config={'evaluation_run_id':run_id,'strategy_name':a.strategy_name,'strategy_version':a.strategy_version,'archetype':a.archetype,'evaluation_mode':a.evaluation_mode,'authoritative_exit':a.authoritative_exit,'policy':cfg,'input_dir':str(a.input_dir.resolve())}; (out/f'{a.strategy_name}_Evaluation_Config.json').write_text(json.dumps(config,indent=2))
 schema={'trade_csv':{'required':['evaluation_run_id','strategy_name','strategy_version','evaluation_mode'],'source_columns':list(trades.columns)},'workbook_sheets':SHEETS,'validation_states':['VALID_AND_ROBUST','VALID_BUT_CONDITIONAL','EXPERIMENTAL','WEAK_EVIDENCE','INVALID_DUE_TO_RISK','INVALID_DUE_TO_NEGATIVE_OOS_ECONOMICS','NOT_SCORABLE_DATA_FAILURE','NOT_SCORABLE_METHOD_FAILURE']}; (out/f'{a.strategy_name}_Evaluation_Schema.json').write_text(json.dumps(schema,indent=2))
 risks.to_csv(out/f'{a.strategy_name}_Risk_Register.csv',index=False); findings=f'# {a.strategy_name} Universal Evaluation Findings\n\n- State: **{state}**\n- Trades: {len(trades):,}; symbols: {trades.symbol.nunique() if "symbol" in trades else "unknown"}.\n- Open liabilities are included in net-result summaries.\n- MFE and ladder touches are diagnostics, not realised profit.\n- Numeric quality score is blocked until all hard gates pass, especially chronological out-of-sample evidence and finite-capital evidence where portfolio returns are claimed.\n'; (out/f'{a.strategy_name}_Evaluation_Findings.md').write_text(findings)
 from docx import Document
 report=Document(); report.add_heading(f'{a.strategy_name} Universal Strategy Evaluation',0); report.add_paragraph(f'Validation state: {state}'); report.add_paragraph(f'Strategy version: {a.strategy_version}; archetype: {a.archetype}; evaluation mode: {a.evaluation_mode}.'); report.add_heading('Executive findings',level=1)
 for row in summary.to_dict('records'): report.add_paragraph(f"{row['metric']}: {row['value']} {row['unit']} ({row['status']})",style='List Bullet')
 report.add_heading('Validation gates',level=1)
 for row in gates.to_dict('records'): report.add_paragraph(f"{row['gate_name']}: {row['status']} — {row['evidence']}",style='List Bullet')
 report.add_heading('Interpretation controls',level=1); report.add_paragraph('Reward/adverse ladders and MFE are path diagnostics. They are not realised return. Portfolio return is valid only when chronological finite-capital allocation and daily equity evidence exist. Open mark-to-market liabilities remain included.'); report.save(out/f'{a.strategy_name}_Evaluation_Report.docx')
 import matplotlib.pyplot as plt
 if len(annual) and 'entry_year' in annual:
  fig,ax=plt.subplots(figsize=(10,5)); ax.bar(annual.entry_year.astype(str),annual.net_pnl); ax.set(title='Annual net result including open MTM',xlabel='Entry year',ylabel='INR'); ax.tick_params(axis='x',rotation=45); fig.tight_layout(); fig.savefig(out/f'{a.strategy_name}_Annual_Net_Result.png',dpi=150); plt.close(fig)
 if 'mfe_pct' in trades and 'mae_pct' in trades:
  fig,ax=plt.subplots(figsize=(8,6)); ax.scatter(trades.mae_pct,trades.mfe_pct,s=16,alpha=.6); ax.axvline(0,color='black',lw=.7); ax.axhline(0,color='black',lw=.7); ax.set(title='Trade path: MFE versus MAE',xlabel='MAE %',ylabel='MFE %'); fig.tight_layout(); fig.savefig(out/f'{a.strategy_name}_MFE_MAE.png',dpi=150); plt.close(fig)
 files=[q for q in out.iterdir() if q.is_file()]; manifest=[{'name':q.name,'path':str(q.resolve()),'sha256':sha(q),'size_bytes':q.stat().st_size} for q in files]; (out/f'{a.strategy_name}_Evidence_Index.json').write_text(json.dumps(manifest,indent=2)); files=[q for q in out.iterdir() if q.is_file()]; sums='\n'.join(f'{sha(q)}  {q.name}' for q in sorted(files))+'\n'; (out/f'{a.strategy_name}_SHA256SUMS.txt').write_text(sums)
 package=out/f'{a.strategy_name}_Evaluation_Package.zip'
 with zipfile.ZipFile(package,'w',zipfile.ZIP_DEFLATED,compresslevel=6) as z:
  for q in out.iterdir():
   if q.is_file() and q!=package:z.write(q,q.name)
 if a.database_url:
  import psycopg
  input_hash=sha(a.input_dir/'trades.csv'); summary_payload={str(r['metric']):r['value'] for r in summary.to_dict('records')}
  with psycopg.connect(a.database_url) as conn:
   conn.execute(DDL.read_text()); conn.execute('''INSERT INTO strategy_eval.universal_evaluation_run (evaluation_run_id,strategy_name,strategy_version,strategy_archetype,evaluation_mode,policy_id,input_path,output_path,actual_start,actual_end,trade_count,validation_state,overall_score,config_json,summary_json,input_sha256) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s::jsonb,%s::jsonb,%s) ON CONFLICT(evaluation_run_id) DO NOTHING''',(run_id,a.strategy_name,a.strategy_version,a.archetype,a.evaluation_mode,cfg['policy_id'],str(a.input_dir.resolve()),str(out.resolve()),trades.entry_date.min().date() if 'entry_date' in trades and trades.entry_date.notna().any() else None,trades.entry_date.max().date() if 'entry_date' in trades and trades.entry_date.notna().any() else None,len(trades),state,None,json.dumps(config),json.dumps(summary_payload,default=str),input_hash))
   with conn.cursor() as cur:
    cur.executemany('INSERT INTO strategy_eval.universal_validation_gate VALUES (%s,%s,%s,%s) ON CONFLICT DO NOTHING',[(run_id,r.gate_name,r.status,r.evidence) for r in gates.itertuples()]); cur.executemany('INSERT INTO strategy_eval.universal_risk_register VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s) ON CONFLICT DO NOTHING',[(run_id,*r) for r in risks.itertuples(index=False,name=None)]); cur.executemany('INSERT INTO strategy_eval.universal_artifact_manifest VALUES (%s,%s,%s,%s,%s) ON CONFLICT DO NOTHING',[(run_id,q.name,str(q.resolve()),sha(q),q.stat().st_size) for q in out.iterdir() if q.is_file()])
 print(json.dumps({'evaluation_run_id':run_id,'state':state,'trades':len(trades),'excel':str(xlsx),'trade_csv':str(trade_path),'package':str(package)},indent=2))
if __name__=='__main__':main()
